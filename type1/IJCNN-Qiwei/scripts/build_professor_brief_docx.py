from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
FIGURE = ROOT / "figures" / "qiwei_research_style_transformer_wm_for_doc.png"
OUTPUT = ROOT / "Qiwei_Model_Brief_for_Professor.docx"


def font(size, bold=False, italic=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else None,
        "/System/Library/Fonts/Supplemental/Arial Italic.ttf" if italic else None,
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def text_center(draw, xy, text, fnt, fill=(20, 20, 20)):
    x, y = xy
    box = draw.textbbox((0, 0), text, font=fnt)
    draw.text((x - (box[2] - box[0]) / 2, y - (box[3] - box[1]) / 2), text, font=fnt, fill=fill)


def rounded(draw, box, fill, outline=None, width=2, radius=14):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill=(31, 35, 40), width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(y2 - y1) >= abs(x2 - x1):
        direction = 1 if y2 > y1 else -1
        pts = [(x2, y2), (x2 - 13, y2 - 24 * direction), (x2 + 13, y2 - 24 * direction)]
    else:
        direction = 1 if x2 > x1 else -1
        pts = [(x2, y2), (x2 - 24 * direction, y2 - 13), (x2 - 24 * direction, y2 + 13)]
    draw.polygon(pts, fill=fill)


def draw_grid(draw, x, y, cell=34):
    for r in range(6):
        for c in range(6):
            color = (218, 218, 218)
            if (c <= r <= c + 2 and c < 3) or (r in (1, 2) and c in (0, 1)):
                color = (169, 211, 184)
            if r >= 3 and c >= 2 and c <= r:
                color = (198, 182, 221)
            draw.rectangle([x + c * cell, y + r * cell, x + (c + 1) * cell, y + (r + 1) * cell], fill=color, outline=(32, 32, 32), width=2)


def draw_small_block(draw, x, y, color, curve_color):
    cell = 28
    for r in range(3):
        for c in range(3):
            draw.rectangle([x + c * cell, y + r * cell, x + (c + 1) * cell, y + (r + 1) * cell], fill=(217, 200, 239), outline=(32, 32, 32), width=2)
    draw.arc([x + 16, y + 10, x + 104, y + 82], start=205, end=335, fill=curve_color, width=5)
    draw.polygon([(x + 96, y + 45), (x + 118, y + 34), (x + 111, y + 59)], fill=curve_color)


def build_figure_png():
    FIGURE.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 980), "white")
    d = ImageDraw.Draw(img)

    title_f = font(40, bold=True)
    sub_f = font(23, italic=True)
    label_f = font(31, bold=True)
    mid_f = font(25)
    small_f = font(19)
    tiny_f = font(16)

    text_center(d, (800, 45), "SSM-Enhanced Transformer Reasoner", title_f)
    text_center(d, (800, 84), "Transformer backbone with latent world-model state injection", sub_f, (82, 82, 82))

    rounded(d, [70, 120, 630, 760], (245, 247, 250), radius=28)
    x0, w, cx = 205, 300, 355
    blocks = [
        ("Scale", (175, 225), (190, 220, 241)),
        ("SwiGLU FFN", (248, 298), (191, 224, 213)),
        ("RMSNorm", (321, 371), (190, 220, 241)),
        ("Semantic Local\nAttention", (420, 492), (245, 200, 172)),
        ("RMSNorm", (515, 565), (190, 220, 241)),
        ("Scale", (625, 675), (190, 220, 241)),
        ("Block-Wise SSM\nWorld Model", (700, 770), (241, 179, 209)),
    ]
    arrow(d, (cx, 180), (cx, 130))
    d.line([(cx, 165), (cx, 780)], fill=(31, 35, 40), width=5)
    for text, (y1, y2), color in blocks:
        rounded(d, [x0, y1, x0 + w, y2], color, radius=10)
        lines = text.split("\n")
        if len(lines) == 1:
            text_center(d, (cx, (y1 + y2) / 2), lines[0], mid_f)
        else:
            text_center(d, (cx, (y1 + y2) / 2 - 15), lines[0], mid_f)
            text_center(d, (cx, (y1 + y2) / 2 + 17), lines[1], mid_f)
    for y in (398, 595):
        d.ellipse([cx - 22, y - 22, cx + 22, y + 22], fill="white", outline=(31, 35, 40), width=4)
        d.line([(cx - 13, y), (cx + 13, y)], fill=(31, 35, 40), width=4)
        d.line([(cx, y - 13), (cx, y + 13)], fill=(31, 35, 40), width=4)
    d.text((535, 204), "x N", font=label_f, fill=(20, 20, 20))
    d.arc([490, 220, 620, 815], start=270, end=92, fill=(31, 35, 40), width=4)
    arrow(d, (525, 456), (505, 456), width=4)
    arrow(d, (525, 735), (505, 735), width=4)

    rounded(d, [755, 135, 1500, 465], (255, 247, 241), radius=28)
    text_center(d, (1128, 195), "Semantic Local Attention", label_f)
    text_center(d, (1128, 232), "nearby premise-answer token matching", sub_f, (82, 82, 82))
    draw_grid(d, 990, 270, 38)
    d.arc([1140, 510, 1270, 575], start=25, end=155, fill=(31, 35, 40), width=4)
    text_center(d, (1205, 590), "k", mid_f)

    rounded(d, [755, 525, 1500, 855], (248, 243, 251), radius=28)
    text_center(d, (1128, 585), "Block-Wise SSM", label_f)
    text_center(d, (1128, 622), "latent scans over token blocks", sub_f, (82, 82, 82))
    draw_small_block(d, 930, 680, (217, 200, 239), (210, 79, 145))
    d.ellipse([870, 705, 908, 743], fill="white", outline=(210, 79, 145), width=5)
    text_center(d, (889, 725), "1", small_f, (210, 79, 145))
    draw_small_block(d, 1210, 680, (217, 200, 239), (37, 139, 197))
    d.ellipse([1150, 705, 1188, 743], fill="white", outline=(37, 139, 197), width=5)
    text_center(d, (1169, 725), "2", small_f, (37, 139, 197))
    draw_small_block(d, 930, 810, (217, 200, 239), (67, 170, 91))
    draw_small_block(d, 1210, 810, (217, 200, 239), (225, 138, 50))
    d.arc([835, 765, 915, 900], start=105, end=255, fill=(31, 35, 40), width=4)
    d.text((800, 820), "b", font=mid_f, fill=(31, 35, 40))
    d.text((817, 835), "h", font=tiny_f, fill=(31, 35, 40))
    d.arc([1000, 900, 1340, 965], start=25, end=155, fill=(31, 35, 40), width=4)
    d.text((1155, 950), "b", font=mid_f, fill=(31, 35, 40))
    d.text((1172, 965), "w", font=tiny_f, fill=(31, 35, 40))
    d.arc([1430, 665, 1510, 850], start=285, end=75, fill=(31, 35, 40), width=4)
    d.text((1518, 755), "T", font=mid_f, fill=(31, 35, 40))

    rounded(d, [95, 865, 210, 915], (190, 220, 241), radius=6)
    text_center(d, (152, 890), "Input", small_f)
    d.polygon([(250, 818), (310, 844), (310, 910), (250, 936)], fill=(191, 224, 213), outline=(31, 35, 40))
    d.text((275, 850), "Encoder", font=small_f, fill=(20, 20, 20))
    rounded(d, [350, 835, 500, 915], (217, 200, 239), outline=(31, 35, 40), radius=3)
    for gx in (400, 450):
        d.line([(gx, 835), (gx, 915)], fill=(31, 35, 40), width=2)
    for gy in (862, 889):
        d.line([(350, gy), (500, gy)], fill=(31, 35, 40), width=2)
    rounded(d, [505, 822, 630, 862], (248, 231, 165), radius=7)
    text_center(d, (567, 842), "BGE-RAG", small_f)
    rounded(d, [505, 875, 630, 915], (190, 220, 241), radius=7)
    text_center(d, (567, 895), "Adapter", small_f)
    rounded(d, [95, 928, 550, 970], "white", outline=(31, 35, 40), radius=10)
    text_center(d, (322, 949), "SSM state injection: K = Wk h + Wks s", small_f)

    img.save(FIGURE, quality=95)


def set_run_font(run, size_pt=None, bold=False, color=None):
    run.font.name = "Arial"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def main():
    build_figure_png()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("SSM-Enhanced Transformer Reasoner")
    set_run_font(run, 18, bold=True, color=(17, 17, 17))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("A shared architecture for Type 1 natural-language logic and Type 2 symbolic logic")
    set_run_font(run, 10.5, color=(80, 80, 80))

    figure_para = doc.add_paragraph()
    figure_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_para.paragraph_format.space_after = Pt(6)
    figure_para.add_run().add_picture(str(FIGURE), width=Inches(5.95))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(10)
    run = caption.add_run("Figure 1. SSM-enhanced Transformer with semantic memory and latent world-model state injection.")
    set_run_font(run, 9, color=(85, 85, 85))

    paragraph_text = (
        "This project proposes an SSM-enhanced Transformer architecture for logical reasoning over educational "
        "physics-style questions. The model keeps a traditional Transformer as the main decision-making backbone, "
        "but improves it with three reasoning-oriented components: a type-specific input adapter, BGE-RAG semantic "
        "memory, and an SSM-based world model. Type 1 questions are natural-language logic problems, where the "
        "adapter extracts premises, answer candidates, semantic evidence, and question categories. Type 2 questions "
        "are symbolic or formal-logic problems, where the adapter extracts predicates, quantifiers, implications, "
        "negations, and contradiction structure. After this front-end difference, both types share the same reasoning "
        "backbone. The SSM world model scans premise-answer token blocks and produces a latent reasoning state, which "
        "is injected into the Transformer attention key as K = W_k h + W_ks s. Therefore, the Transformer attends not "
        "only to token semantics, but also to inferred premise-to-answer reasoning flow. The purpose of this design is "
        "to move beyond simple text matching and give the model a more explicit internal mechanism for support, "
        "contradiction, uncertainty, and candidate comparison."
    )
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing = 1.12
    body.paragraph_format.space_after = Pt(0)
    run = body.add_run(paragraph_text)
    set_run_font(run, 10.5, color=(20, 20, 20))

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
